import os
import openpyxl
import docx
import pdfplumber
import re
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from openpyxl import Workbook
from openpyxl.cell import Cell
from openpyxl.worksheet.worksheet import Worksheet

# 资源路径
resourcesPath = "./resources"

# 所有文件的完整路径（由于文件分布在不同目录的文件夹下，必须用os.walk）
itemPaths = []
for root, dir, items in os.walk(top=resourcesPath):
    itemPaths.extend(
        [os.path.join(root, item) for item in items]
    )

# 电话号码的正则模式（提前编译好）
phonePattern = re.compile(pattern=r"1[3|4|5|7|8][0-9]{9}")


# 读取word文档的手机号
def get_phones_from_docx(docx_path: str) -> list:
    phone_list = []
    # 打开文件
    word_file = docx.Document(docx=docx_path)  # type: Document
    # 读取正文
    for para in word_file.paragraphs:
        phone_list.extend(
            match_phone_list(value=para.text)
        )
    # 读取表格
    for table in word_file.tables:
        for row in table.rows:
            for cell in row.cells:
                phone_list.extend(
                    match_phone_list(value=cell.text)
                )

    return phone_list


# 读取excel文档的手机号
def get_phones_from_xlsx(xlsx_path: str) -> list:
    # test: 打印excel表格路径
    # print(xlsx_path)
    phone_list = []
    # 打开文件
    work_book = openpyxl.load_workbook(filename=xlsx_path)  # type: Workbook
    # 逐行遍历，匹配电话号码
    for sheet_name in work_book.sheetnames:
        # 选工作表
        work_sheet = work_book[sheet_name]  # type: Worksheet
        # 逐行遍历
        for rowData in work_sheet.iter_rows(values_only=True):
            # test: 打印当前行的数据
            # print(rowData)
            for cell in rowData:
                # test: 打印单元格中的内容
                # print(cell)
                if cell is not None and cell != "":
                    phone_numbers = match_phone_list(value=cell)
                    # test: 打印手机号查询的结果
                    # if phone_numbers is not None and len(phone_numbers) > 0:
                    #     print("找到的手机号：")
                    #     print(phone_numbers)
                    phone_list.extend(phone_numbers)

    return phone_list


# 读取pdf文档中的手机号
def get_phones_from_pdf(pdf_path: str) -> list:
    phone_list = []
    # 打开文件
    pdf = pdfplumber.open(path_or_fp=pdf_path)
    # 逐页遍历
    for page in pdf.pages:
        # 提取正文文本内容
        phone_list.extend(
            match_phone_list(
                value=page.extract_text()
            )
        )
        # 读取表格内容
        for table in page.extract_tables():
            for row in table:
                for rowData in row:
                    phone_list.extend(
                        match_phone_list(
                            value=rowData
                        )
                    )

    return phone_list


# 取出所有的电话号码
def get_phone_list() -> list:
    phone_list = []
    for itemPath in itemPaths:
        # 扩展名
        extension = os.path.splitext(p=itemPath)[-1].lower()
        # 查找word文档
        if extension == ".docx":
            phone_list.extend(
                get_phones_from_docx(docx_path=itemPath)
            )
        # 查找excel文档
        elif extension == ".xlsx":
            phone_list.extend(
                get_phones_from_xlsx(xlsx_path=itemPath)
            )
        # 查找pdf文档
        elif extension == ".pdf":
            phone_list.extend(
                get_phones_from_pdf(pdf_path=itemPath)
            )
        else:
            continue

    return phone_list


def match_phone_list(value):
    result = phonePattern.findall(string=str(value))
    return result


phone_list = list(set(get_phone_list()))
print(phone_list)
# test: 输出长度进行测试
# print(len(phone_list))

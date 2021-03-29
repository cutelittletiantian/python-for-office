import os
import shutil

import docx
import openpyxl

from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from openpyxl import Workbook
from openpyxl.cell import Cell
from openpyxl.worksheet.worksheet import Worksheet


# 设置工作区
os.chdir(path="./resources/exer_2")
# 原文路径
articlePath = "article"
# 汇总表路径
summaryPath = "作者投稿-20年.xlsx"

# 修改后的文档的保存路径
modifyPath = "modify"
if not os.path.exists(path=modifyPath):
    os.mkdir(path=modifyPath)

# 提取所有word文档中的所有重要信息[文章题目，姓名，学院，邮箱]
wordInfo = []
# 获取所有word文档所在的列表
articleList = os.listdir(path=articlePath)
# 扫描所有文件
for articleItem in articleList:
    # 组装完整路径
    articleItemPath = os.path.join(articlePath, articleItem)

    # 读取文件
    if os.path.splitext(p=articleItem)[1] not in [".docx", ".doc"]:
        continue
    articleDoc = docx.Document(docx=articleItemPath)  # type: Document

    # 取出[文章题目，姓名，学院，邮箱]信息
    articleInfo = []
    # 文章题目
    articleInfo.append(articleDoc.paragraphs[0].text)
    # 姓名、学院、邮箱
    articleInfo += articleDoc.paragraphs[2].text.split(sep=" ")
    # 当前信息添加到wordInfo中
    wordInfo.append(articleInfo)

# print(wordInfo)

# 扫描所有汇总表
excelInfo = dict()
# 加载excel文件
summaryBook = openpyxl.load_workbook(filename=summaryPath)  # type: Workbook
# 逐个扫描工作表
for sheetName in summaryBook.sheetnames:
    excelInfo[sheetName] = []
    # 打开工作表
    summarySheet = summaryBook[sheetName]  # type: Worksheet
    # 跳过表头，逐行读入数据
    for rowData in summarySheet.iter_rows(min_row=2, values_only=True):
        # 录入信息
        excelInfo[sheetName].append(list(rowData))
# 关闭汇总表，节省资源
summaryBook.close()

# print(excelInfo)


# 匹配文档，按格式改名
for articleInfo in wordInfo:
    # 扫描excel信息词典
    for month, content in excelInfo.items():
        if articleInfo in content:
            # 匹配到，执行改名
            name = articleInfo[1]  # 姓名
            # 按照格式重命名文件，将修改后的文件拷贝到dst所在目录下
            # 重命名+拷贝：copy，取什么名字，dst下面的文件就重命名成什么样
            shutil.copy(src=os.path.join(articlePath, f"{name}投稿.docx"),
                        dst=os.path.join(modifyPath, f"2020年{month}{name}投稿.docx")
                        )
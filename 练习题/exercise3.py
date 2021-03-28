import openpyxl
import docx
import os

from docx.document import Document
from openpyxl import Workbook
from openpyxl.worksheet.worksheet import Worksheet

# 工作路径
os.chdir(path="./resources/exer_3")
# 来稿路径
articlePath = "sub"
# 汇总表
summaryBookPath = "作者投稿.xlsx"
# 工作表名
workSheetName = "6月"

# 加载工作簿
summaryBook = openpyxl.load_workbook(filename=summaryBookPath)  # type: Workbook
# 读取工作表
summarySheet = summaryBook[workSheetName]  # type: Worksheet

# 获取投稿文件列表
articleList = os.listdir(path=articlePath)
# 扫描逐个处理
for articleItem in articleList:
    # 上来不忘组装完整路径
    articleItemPath = os.path.join(articlePath, articleItem)
    # 加载文档
    articleDoc = docx.Document(docx=articleItemPath)  # type: Document
    # 提取所需的信息[文章题目，姓名，地址，邮箱]
    articleInfo = []
    # 取出标题
    articleInfo.append(articleDoc.paragraphs[0].text)
    # 取出姓名、地址、邮箱
    articleInfo += articleDoc.paragraphs[2].text.split(" ")
    # 添加信息到工作表一行中
    summarySheet.append(articleInfo)

# 保存工作表
summaryBook.save(filename=summaryBookPath)
summaryBook.close()
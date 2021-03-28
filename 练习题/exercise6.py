import os
import docx
from docx.document import Document

# 设置工作区
os.chdir(path="./resources/exer_6")

# 原文件路径
sourcePath = "stu"
# 目标文件的路径
targetPath = "modify"
# 如果目标路径不存在，先创建好
if not os.path.exists(targetPath):
    os.mkdir(path=targetPath)

# 读取待修改文件列表
proofList = os.listdir(path=sourcePath)
for proofItem in proofList:
    # 健壮性判断：是否是word文档
    extension = os.path.splitext(proofItem)[1]
    if extension.lower() not in [".docx", ".doc"]:
        continue

    # 组装完整路径
    proofItemPath = os.path.join(sourcePath, proofItem)
    # 打开相应文档
    proofDoc = docx.Document(docx=proofItemPath)  # type: Document
    # 修改”夜曲大学“为”夜曲大学研究生院“
    # proofDoc.paragraphs[7].add_run(text="研究生院")
    proofDoc.paragraphs[7].text = "夜曲大学研究生院"

    # 组装保存文档路径
    targetItemPath = os.path.join(targetPath, proofItem)
    # 保存
    proofDoc.save(path_or_stream=targetItemPath)
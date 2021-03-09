import os
import shutil
import docx
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# 偏好表路径
preferenceListPath = "./resources/ex5_5/调查表"

# 获取意向文档列表（这一步必须先于创建文件夹，否则文件夹也会算进来）
allPreferences = os.listdir(path=preferenceListPath)

# 预先创建好没创建的文件夹
for gender in ["男", "女"]:
    genderPath = os.path.join(preferenceListPath, gender)
    if not os.path.exists(genderPath):
        os.mkdir(genderPath)
    for roomType in ["需要安静的学习环境", "不太需要安静的学习环境"]:
        roomTypePath = os.path.join(genderPath, roomType)
        if not os.path.exists(roomTypePath):
            os.mkdir(roomTypePath)

# 遍历所有意向文档
for preference in allPreferences:
    # 组装完整路径
    preferencePath = os.path.join(preferenceListPath, preference)
    # 载入文档
    doc = docx.Document(docx=preferencePath)  # type: Document
    # 提取性别
    sex = doc.paragraphs[2].runs[1].text
    # 获取安静环境得分
    score = int(doc.tables[0].cell(row_idx=2, col_idx=1).text)

    # 判断移动
    targetPath = ""
    if score > 5:
        targetPath = os.path.join(preferenceListPath, sex, "需要安静的学习环境")
    else:
        targetPath = os.path.join(preferenceListPath, sex, "不太需要安静的学习环境")

    # 执行移动
    shutil.move(src=preferencePath, dst=targetPath)
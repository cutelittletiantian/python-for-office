import os
import docx
import openpyxl
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from openpyxl import Workbook
from openpyxl.worksheet.worksheet import Worksheet

schoolListPath = "./resources/ex5_6/统计表"
gatherPath = "./resources/ex5_6/我是健康小卫士打分表.xlsx"

# 加载excel表格
gatherBook = openpyxl.load_workbook(filename=gatherPath)  # type: Workbook
# 加载工作表
gatherSheet = gatherBook["我是健康小卫士"]   # type: Worksheet

# 读word文档
schoolItems = os.listdir(path=schoolListPath)
for school in schoolItems:
    # 组装完整路径
    schoolPath = os.path.join(schoolListPath, school)
    # 打开文档
    schoolDoc = docx.Document(docx=schoolPath)  # type: Document

    # 从文档中提取数据信息
    rowContent = ["", "", "", ""]

    # 读学校名
    schoolName = schoolDoc.paragraphs[1].runs[1].text
    rowContent[0] = schoolName

    # 提取表中表格内容
    schoolTable = schoolDoc.tables[0]  # type: Table
    row_index = 2 # 我是健康小卫士所在行号
    # 读一行
    for col_index in [2, 3, 4]:
        if schoolTable.cell(row_idx=row_index, col_idx=col_index).text == "":
            break
        else:
            rowContent[col_index - 1] = schoolTable.cell(row_idx=row_index, col_idx=col_index).text
    else:
        # for循环完整执行后，表明一行数据都完整了，存一行
        gatherSheet.append(iterable=rowContent)

# 保存表格文件
gatherBook.save(filename=gatherPath)
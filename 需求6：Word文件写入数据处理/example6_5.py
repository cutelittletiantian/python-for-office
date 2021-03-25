import os
import openpyxl
import docx
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from openpyxl import Workbook
from openpyxl.worksheet.worksheet import Worksheet

# 工作区设置
os.chdir("./resources/ex6_5/")

# 模板路径
templatePath = "合同模版.docx"
# 采购信息表路径
purchasePath = "采购信息列表.xlsx"
# 目标保存文件夹
targetPath = "合同"
# 如果没有合同文件夹就创建
if not os.path.exists(targetPath):
    os.mkdir(path=targetPath)


#####################################################################
# 文档中的字符串替换函数
def replace_doc_info(doc: Document, org_str: str, dst_str: str):
    # 扫描段落及样式块
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # 执行子串替换
            run.text = run.text.replace(org_str, dst_str)

    # 表格执行替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 执行子串替换
                cell.text = cell.text.replace(org_str, dst_str)
#####################################################################


# 读工作簿+读工作表
purchaseSheet = openpyxl.load_workbook(filename=purchasePath, data_only=True)["6月采购"]  # type: Worksheet

# 获取表头字符串
# ('【合同编号】', '【甲方】', '【乙方】', '【货物名称】', '【单价】', '【数量】',
# '【商品总价】', '【总税额】', '【总金额】', '【总金额（大写）】')
titleData = []
for rowData in purchaseSheet.iter_rows(min_row=1, max_row=1, values_only=True):
    titleData = rowData

# 遍历工作表各行（跳过表头），例如：
# ('CG01200601', '开网店的阿珍', '开口哭牌的供应商', '开口哭牌蟹黄味薯片', 15, 500,
# 7500, 1275, 8775, 8775)
# 对每一种产品生成对应的合同
for rowData in purchaseSheet.iter_rows(min_row=2, values_only=True):
    # 读模板文件
    templateDoc = docx.Document(docx=templatePath)  # type: Document

    # 扫描所有列（同时获取列号，列号与表头一一对应）
    for columnIndex, content in enumerate(rowData):
        # 将可能的数值类型转为字符串
        content = str(content)
        # 模板中替换表头下对应内容
        replace_doc_info(doc=templateDoc, org_str=titleData[columnIndex], dst_str=content)

    # 单独拎出来产品名称（列下标为3）
    productName = rowData[3]
    # 组装文件完整路径+保存模板
    templateDoc.save(path_or_stream=os.path.join(targetPath, f"采购合同_{productName}.docx"))
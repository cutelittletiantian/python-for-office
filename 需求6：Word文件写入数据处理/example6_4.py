import os

import openpyxl
import docx
from docx.document import Document
from docx.shared import Cm
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from openpyxl import Workbook
from openpyxl.utils import cell
from openpyxl.worksheet.worksheet import Worksheet

# 基本配置参数
# 设置工作路径
os.chdir(path="./resources/ex6_4/volunteer")
# 名单路径
nameListPath = "志愿者名单.xlsx"
# 工牌模版路径
templatePath = "志愿者工牌模版.docx"
# 目标路径
targetPath = "志愿者"
if not os.path.exists(targetPath):
    os.mkdir(path=targetPath)
# 照片所在路径
photoPath = "photo"
# 抽出照片列表
photoItems = os.listdir(path=photoPath)


# 生成工牌函数
def generate_card(name: str, group: str):
    photo_item_path = None
    # 找照片是否存在
    if f"{name}.png" in photoItems:
        # 组装完整路径
        photo_item_path = os.path.join(photoPath, f"{name}.png")
    # 现有参数：姓名、组别、照片完整路径，可以愉快地合成工牌了
    # 打开模板
    templateDoc = docx.Document(docx=templatePath)  # type: Document
    # 添加姓名
    templateDoc.paragraphs[7].runs[0].add_text(name)
    # 添加组别
    templateDoc.paragraphs[8].runs[0].add_text(group)
    # 添加图片
    if photo_item_path is not None:
        cell_paragraph = templateDoc.tables[0].cell(row_idx=0, col_idx=0).paragraphs[0]
        cell_run = cell_paragraph.add_run()
        cell_run.add_picture(image_path_or_stream=photo_item_path, width=docx.shared.Cm(5))
    # 另存
    templateDoc.save(path_or_stream=os.path.join(targetPath, f"工牌_{name}.docx"))


# 导入名单信息
nameSheet = openpyxl.load_workbook(filename=nameListPath)["名单"]  # type: Worksheet
# 遍历每行工作人员信息
for rowData in nameSheet.iter_rows(min_row=2):
    name = rowData[0].value
    group = rowData[1].value
    # 完成一个工牌，生成一个在"志愿者"文件夹下的Word文档
    # 命名为"工牌_{name}.docx"
    generate_card(name=name, group=group)

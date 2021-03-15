import os
import random
import shutil
import docx
from docx.document import Document

from docx.text.paragraph import Paragraph
from docx.text.run import Run

# 设置工作目录
os.chdir(path="./resources/ex6_3")
# 试卷存储路径，不存在就创建
targetPath = "examination"
if not os.path.exists(path=targetPath):
    os.mkdir(path=targetPath)
# 试卷模板路径
templatePath = "试卷模版.docx"
# 题库路径
quizLibPath = "题库.docx"

# 打开题库
quizLibDoc = docx.Document(docx=quizLibPath)  # type: Document

# 提取第一大题题库的所有内容
partOneLib = [paragraph.text for paragraph in quizLibDoc.paragraphs[0:12]]
# 提取第二大题题库的所有内容
partTwoLib = [paragraph.text for paragraph in quizLibDoc.paragraphs[14:22]]

# 执行72次抽题组卷
for execTime in range(72):
    # 打开模板
    templateDoc = docx.Document(docx=templatePath)  # type: Document
    # 抽第一大题（6道）
    # counts参数在Python 3.9版本才有，表示权重，3.8.x用户请勿使用，避免兼容问题
    partOne = random.sample(population=partOneLib, k=6)
    # 抽第二大题（4道）
    partTwo = random.sample(population=partTwoLib, k=4)
    # 第一大题写入模板
    for index, quiz in enumerate(partOne, start=0):
        templateDoc.paragraphs[index].runs[0].add_text(quiz)
    # 第二大题写入模板
    for index, quiz in enumerate(partTwo, start=6):
        templateDoc.paragraphs[index].runs[0].add_text(quiz)
    # 另存
    templateDoc.save(path_or_stream=os.path.join(targetPath, f"试卷_{execTime+1}.docx"))